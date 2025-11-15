"""
Document generation utilities for creating PDF and DOCX resumes.

This module provides functions to generate professional resume documents
from ResumeModel instances.
"""

from typing import Optional
from pathlib import Path
from datetime import datetime
import io

from jinja2 import Environment, FileSystemLoader, Template
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

from modules.models import ResumeModel, ExperienceItem, EducationItem

# Get the templates directory
TEMPLATES_DIR = Path(__file__).parent.parent / "templates"


def _get_jinja_env() -> Environment:
    """Create and configure Jinja2 environment."""
    if TEMPLATES_DIR.exists():
        return Environment(
            loader=FileSystemLoader(str(TEMPLATES_DIR)),
            autoescape=True,
            trim_blocks=True,
            lstrip_blocks=True
        )
    else:
        # Use a template from string if templates directory doesn't exist
        return Environment(autoescape=True, trim_blocks=True, lstrip_blocks=True)


def _get_inline_template() -> str:
    """Return inline HTML template for resume."""
    return """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resume - {{ resume.name }}</title>
    <style>
        @page {
            size: letter;
            margin: 0.75in;
        }

        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Calibri', 'Arial', sans-serif;
            font-size: 11pt;
            line-height: 1.4;
            color: #333;
            background: white;
        }

        .container {
            max-width: 8.5in;
            margin: 0 auto;
            padding: 0.5in 0;
        }

        /* Header Section */
        .header {
            text-align: center;
            margin-bottom: 20px;
            border-bottom: 2px solid #2c3e50;
            padding-bottom: 15px;
        }

        .name {
            font-size: 24pt;
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 8px;
        }

        .headline {
            font-size: 12pt;
            color: #555;
            font-weight: 600;
            margin-bottom: 10px;
        }

        .contact-info {
            font-size: 10pt;
            color: #666;
            margin: 5px 0;
        }

        .contact-info a {
            color: #2c3e50;
            text-decoration: none;
        }

        .links {
            font-size: 10pt;
            color: #666;
            margin-top: 5px;
        }

        .links a {
            color: #2c3e50;
            text-decoration: none;
            margin: 0 8px;
        }

        /* Section Headings */
        .section {
            margin: 18px 0;
        }

        .section-title {
            font-size: 14pt;
            font-weight: bold;
            color: #2c3e50;
            text-transform: uppercase;
            border-bottom: 1.5px solid #2c3e50;
            padding-bottom: 4px;
            margin-bottom: 12px;
            letter-spacing: 0.5px;
        }

        /* Summary Section */
        .summary-text {
            text-align: justify;
            margin-bottom: 10px;
            line-height: 1.5;
        }

        /* Experience Section */
        .experience-item {
            margin-bottom: 16px;
            page-break-inside: avoid;
        }

        .experience-header {
            display: flex;
            justify-content: space-between;
            align-items: baseline;
            margin-bottom: 4px;
        }

        .job-title {
            font-size: 12pt;
            font-weight: bold;
            color: #2c3e50;
        }

        .company-name {
            font-size: 11pt;
            font-weight: 600;
            color: #555;
        }

        .job-dates {
            font-size: 10pt;
            color: #666;
            font-style: italic;
            white-space: nowrap;
        }

        .job-location {
            font-size: 10pt;
            color: #666;
            margin-bottom: 6px;
        }

        .bullets {
            list-style-type: disc;
            margin-left: 20px;
            margin-top: 6px;
        }

        .bullets li {
            margin-bottom: 4px;
            line-height: 1.4;
        }

        .experience-skills {
            font-size: 10pt;
            color: #555;
            margin-top: 6px;
            font-style: italic;
        }

        /* Education Section */
        .education-item {
            margin-bottom: 12px;
        }

        .degree {
            font-size: 11.5pt;
            font-weight: bold;
            color: #2c3e50;
        }

        .institution {
            font-size: 11pt;
            color: #555;
            font-weight: 600;
        }

        .education-details {
            font-size: 10pt;
            color: #666;
            margin-top: 3px;
        }

        /* Skills Section */
        .skills-list {
            display: flex;
            flex-wrap: wrap;
            gap: 8px;
        }

        .skill-item {
            background: #f0f0f0;
            padding: 4px 10px;
            border-radius: 3px;
            font-size: 10pt;
            color: #333;
        }

        /* Additional Sections */
        .simple-list {
            list-style-type: disc;
            margin-left: 20px;
        }

        .simple-list li {
            margin-bottom: 4px;
            line-height: 1.3;
        }

        .inline-list {
            font-size: 10.5pt;
            line-height: 1.5;
        }

        /* Print Optimization */
        @media print {
            body {
                margin: 0;
                padding: 0;
            }

            .container {
                padding: 0;
            }

            .section {
                page-break-inside: avoid;
            }
        }
    </style>
</head>
<body>
    <div class="container">
        <!-- Header -->
        <div class="header">
            {% if resume.name %}
            <div class="name">{{ resume.name }}</div>
            {% endif %}

            {% if resume.headline %}
            <div class="headline">{{ resume.headline }}</div>
            {% endif %}

            <div class="contact-info">
                {% set contact_parts = [] %}
                {% if resume.email %}{% set _ = contact_parts.append(resume.email) %}{% endif %}
                {% if resume.phone %}{% set _ = contact_parts.append(resume.phone) %}{% endif %}
                {% if resume.location %}{% set _ = contact_parts.append(resume.location) %}{% endif %}
                {{ contact_parts|join(' • ') }}
            </div>

            {% if resume.linkedin or resume.github or resume.portfolio %}
            <div class="links">
                {% if resume.linkedin %}<a href="{{ resume.linkedin }}">LinkedIn</a>{% endif %}
                {% if resume.github %}<a href="{{ resume.github }}">GitHub</a>{% endif %}
                {% if resume.portfolio %}<a href="{{ resume.portfolio }}">Portfolio</a>{% endif %}
            </div>
            {% endif %}
        </div>

        <!-- Professional Summary -->
        {% if resume.summary %}
        <div class="section">
            <div class="section-title">Professional Summary</div>
            <div class="summary-text">{{ resume.summary }}</div>
        </div>
        {% endif %}

        <!-- Professional Experience -->
        {% if resume.experiences %}
        <div class="section">
            <div class="section-title">Professional Experience</div>
            {% for exp in resume.experiences %}
            <div class="experience-item">
                <div class="experience-header">
                    <div>
                        <span class="job-title">{{ exp.title }}</span>
                        {% if exp.company %}
                        <span> • </span>
                        <span class="company-name">{{ exp.company }}</span>
                        {% endif %}
                    </div>
                    <div class="job-dates">
                        {% if exp.start_date or exp.end_date %}
                        {{ exp.start_date or 'N/A' }} – {{ exp.end_date or 'Present' }}
                        {% endif %}
                    </div>
                </div>

                {% if exp.location %}
                <div class="job-location">{{ exp.location }}</div>
                {% endif %}

                {% if exp.bullets %}
                <ul class="bullets">
                    {% for bullet in exp.bullets %}
                    <li>{{ bullet }}</li>
                    {% endfor %}
                </ul>
                {% endif %}

                {% if exp.skills %}
                <div class="experience-skills">
                    <strong>Skills:</strong> {{ exp.skills|join(', ') }}
                </div>
                {% endif %}
            </div>
            {% endfor %}
        </div>
        {% endif %}

        <!-- Skills -->
        {% if resume.skills %}
        <div class="section">
            <div class="section-title">Skills</div>
            <div class="skills-list">
                {% for skill in resume.skills %}
                <span class="skill-item">{{ skill }}</span>
                {% endfor %}
            </div>
        </div>
        {% endif %}

        <!-- Education -->
        {% if resume.education %}
        <div class="section">
            <div class="section-title">Education</div>
            {% for edu in resume.education %}
            <div class="education-item">
                <div>
                    <span class="degree">{{ edu.degree }}</span>
                    {% if edu.field_of_study %}
                    <span> in {{ edu.field_of_study }}</span>
                    {% endif %}
                </div>
                {% if edu.institution %}
                <div class="institution">{{ edu.institution }}</div>
                {% endif %}
                <div class="education-details">
                    {% set edu_parts = [] %}
                    {% if edu.graduation_date %}{% set _ = edu_parts.append('Graduated: ' + edu.graduation_date) %}{% endif %}
                    {% if edu.gpa %}{% set _ = edu_parts.append('GPA: ' + edu.gpa) %}{% endif %}
                    {{ edu_parts|join(' • ') }}
                </div>
                {% if edu.honors %}
                <div class="education-details">
                    <strong>Honors:</strong> {{ edu.honors|join(', ') }}
                </div>
                {% endif %}
            </div>
            {% endfor %}
        </div>
        {% endif %}

        <!-- Certifications -->
        {% if resume.certifications %}
        <div class="section">
            <div class="section-title">Certifications</div>
            <ul class="simple-list">
                {% for cert in resume.certifications %}
                <li>{{ cert }}</li>
                {% endfor %}
            </ul>
        </div>
        {% endif %}

        <!-- Projects -->
        {% if resume.projects %}
        <div class="section">
            <div class="section-title">Projects</div>
            <ul class="simple-list">
                {% for project in resume.projects %}
                <li>
                    {% if project is mapping %}
                        {% if project.name %}<strong>{{ project.name }}:</strong> {% endif %}
                        {{ project.description or project.get('details', '') }}
                    {% else %}
                        {{ project }}
                    {% endif %}
                </li>
                {% endfor %}
            </ul>
        </div>
        {% endif %}

        <!-- Awards -->
        {% if resume.awards %}
        <div class="section">
            <div class="section-title">Awards & Honors</div>
            <ul class="simple-list">
                {% for award in resume.awards %}
                <li>{{ award }}</li>
                {% endfor %}
            </ul>
        </div>
        {% endif %}

        <!-- Languages -->
        {% if resume.languages %}
        <div class="section">
            <div class="section-title">Languages</div>
            <div class="inline-list">{{ resume.languages|join(', ') }}</div>
        </div>
        {% endif %}
    </div>
</body>
</html>
"""


def generate_html(resume_model: ResumeModel) -> str:
    """
    Generate HTML representation of resume.

    Args:
        resume_model: The resume data to render

    Returns:
        HTML string of the rendered resume
    """
    env = _get_jinja_env()

    # Try to load template from file, fall back to inline template
    try:
        if TEMPLATES_DIR.exists():
            template = env.get_template("resume_template.html")
        else:
            template = env.from_string(_get_inline_template())
    except Exception:
        template = env.from_string(_get_inline_template())

    # Render template with resume data
    html_content = template.render(resume=resume_model)
    return html_content


def generate_pdf(resume_model: ResumeModel) -> bytes:
    """
    Generate PDF representation of resume.

    Args:
        resume_model: The resume data to render

    Returns:
        PDF file as bytes

    Raises:
        ImportError: If weasyprint is not installed
    """
    try:
        from weasyprint import HTML, CSS
    except ImportError:
        raise ImportError(
            "weasyprint is required for PDF generation. "
            "Install it with: pip install weasyprint"
        )

    # Generate HTML first
    html_content = generate_html(resume_model)

    # Convert HTML to PDF
    pdf_file = HTML(string=html_content).write_pdf()

    return pdf_file


def generate_docx(resume_model: ResumeModel) -> bytes:
    """
    Generate DOCX representation of resume using python-docx.

    Args:
        resume_model: The resume data to render

    Returns:
        DOCX file as bytes
    """
    doc = Document()

    # Configure document margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Define custom styles
    styles = doc.styles

    # Heading style for name
    if 'ResumeName' not in styles:
        name_style = styles.add_style('ResumeName', WD_STYLE_TYPE.PARAGRAPH)
        name_style.font.name = 'Calibri'
        name_style.font.size = Pt(24)
        name_style.font.bold = True
        name_style.font.color.rgb = RGBColor(44, 62, 80)
        name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_style.paragraph_format.space_after = Pt(4)

    # Header Section
    if resume_model.name:
        name_para = doc.add_paragraph(resume_model.name, style='ResumeName')

    if resume_model.headline:
        headline = doc.add_paragraph(resume_model.headline)
        headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        headline.runs[0].font.size = Pt(12)
        headline.runs[0].font.color.rgb = RGBColor(85, 85, 85)
        headline.runs[0].bold = True

    # Contact Information
    contact_parts = []
    if resume_model.email:
        contact_parts.append(resume_model.email)
    if resume_model.phone:
        contact_parts.append(resume_model.phone)
    if resume_model.location:
        contact_parts.append(resume_model.location)

    if contact_parts:
        contact_para = doc.add_paragraph(' • '.join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(10)
        contact_para.runs[0].font.color.rgb = RGBColor(102, 102, 102)

    # Links
    link_parts = []
    if resume_model.linkedin:
        link_parts.append(f"LinkedIn: {resume_model.linkedin}")
    if resume_model.github:
        link_parts.append(f"GitHub: {resume_model.github}")
    if resume_model.portfolio:
        link_parts.append(f"Portfolio: {resume_model.portfolio}")

    if link_parts:
        links_para = doc.add_paragraph(' • '.join(link_parts))
        links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links_para.runs[0].font.size = Pt(10)
        links_para.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    # Add horizontal line after header
    doc.add_paragraph('_' * 80).runs[0].font.color.rgb = RGBColor(44, 62, 80)

    # Helper function to add section heading
    def add_section_heading(title: str):
        heading = doc.add_paragraph()
        run = heading.add_run(title.upper())
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(44, 62, 80)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        # Add underline using a separate paragraph
        underline = doc.add_paragraph('_' * 80)
        underline.runs[0].font.size = Pt(8)
        underline.runs[0].font.color.rgb = RGBColor(44, 62, 80)
        underline.paragraph_format.space_after = Pt(8)

    # Professional Summary
    if resume_model.summary:
        add_section_heading('Professional Summary')
        summary_para = doc.add_paragraph(resume_model.summary)
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        summary_para.paragraph_format.space_after = Pt(10)

    # Professional Experience
    if resume_model.experiences:
        add_section_heading('Professional Experience')

        for exp in resume_model.experiences:
            # Job title and company
            exp_header = doc.add_paragraph()
            title_run = exp_header.add_run(exp.title)
            title_run.font.size = Pt(12)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(44, 62, 80)

            if exp.company:
                exp_header.add_run(' • ')
                company_run = exp_header.add_run(exp.company)
                company_run.font.size = Pt(11)
                company_run.font.bold = True
                company_run.font.color.rgb = RGBColor(85, 85, 85)

            # Dates
            if exp.start_date or exp.end_date:
                dates_para = doc.add_paragraph()
                dates_run = dates_para.add_run(
                    f"{exp.start_date or 'N/A'} – {exp.end_date or 'Present'}"
                )
                dates_run.font.size = Pt(10)
                dates_run.font.italic = True
                dates_run.font.color.rgb = RGBColor(102, 102, 102)

            # Location
            if exp.location:
                loc_para = doc.add_paragraph(exp.location)
                loc_para.runs[0].font.size = Pt(10)
                loc_para.runs[0].font.color.rgb = RGBColor(102, 102, 102)
                loc_para.paragraph_format.space_after = Pt(4)

            # Bullets
            for bullet in exp.bullets:
                bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                bullet_para.runs[0].font.size = Pt(11)
                bullet_para.paragraph_format.left_indent = Inches(0.25)
                bullet_para.paragraph_format.space_after = Pt(2)

            # Skills for this experience
            if exp.skills:
                skills_para = doc.add_paragraph()
                skills_para.add_run('Skills: ').font.bold = True
                skills_para.add_run(', '.join(exp.skills))
                skills_para.runs[0].font.size = Pt(10)
                skills_para.runs[1].font.size = Pt(10)
                skills_para.runs[0].font.italic = True
                skills_para.runs[1].font.italic = True
                skills_para.paragraph_format.space_after = Pt(10)
            else:
                doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Skills
    if resume_model.skills:
        add_section_heading('Skills')
        skills_para = doc.add_paragraph(', '.join(resume_model.skills))
        skills_para.runs[0].font.size = Pt(11)
        skills_para.paragraph_format.space_after = Pt(10)

    # Education
    if resume_model.education:
        add_section_heading('Education')

        for edu in resume_model.education:
            # Degree
            edu_para = doc.add_paragraph()
            degree_run = edu_para.add_run(edu.degree)
            degree_run.font.size = Pt(11.5)
            degree_run.font.bold = True
            degree_run.font.color.rgb = RGBColor(44, 62, 80)

            if edu.field_of_study:
                edu_para.add_run(f" in {edu.field_of_study}")
                edu_para.runs[-1].font.size = Pt(11.5)

            # Institution
            if edu.institution:
                inst_para = doc.add_paragraph(edu.institution)
                inst_para.runs[0].font.size = Pt(11)
                inst_para.runs[0].font.bold = True
                inst_para.runs[0].font.color.rgb = RGBColor(85, 85, 85)

            # Details
            details = []
            if edu.graduation_date:
                details.append(f"Graduated: {edu.graduation_date}")
            if edu.gpa:
                details.append(f"GPA: {edu.gpa}")

            if details:
                details_para = doc.add_paragraph(' • '.join(details))
                details_para.runs[0].font.size = Pt(10)
                details_para.runs[0].font.color.rgb = RGBColor(102, 102, 102)

            # Honors
            if edu.honors:
                honors_para = doc.add_paragraph()
                honors_para.add_run('Honors: ').font.bold = True
                honors_para.add_run(', '.join(edu.honors))
                honors_para.runs[0].font.size = Pt(10)
                honors_para.runs[1].font.size = Pt(10)

            doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Certifications
    if resume_model.certifications:
        add_section_heading('Certifications')
        for cert in resume_model.certifications:
            cert_para = doc.add_paragraph(cert, style='List Bullet')
            cert_para.runs[0].font.size = Pt(11)
            cert_para.paragraph_format.left_indent = Inches(0.25)

    # Projects
    if resume_model.projects:
        add_section_heading('Projects')
        for project in resume_model.projects:
            if isinstance(project, dict):
                proj_text = ""
                if project.get('name'):
                    proj_text = f"{project['name']}: "
                proj_text += project.get('description', project.get('details', ''))
                proj_para = doc.add_paragraph(proj_text, style='List Bullet')
            else:
                proj_para = doc.add_paragraph(str(project), style='List Bullet')
            proj_para.runs[0].font.size = Pt(11)
            proj_para.paragraph_format.left_indent = Inches(0.25)

    # Awards
    if resume_model.awards:
        add_section_heading('Awards & Honors')
        for award in resume_model.awards:
            award_para = doc.add_paragraph(award, style='List Bullet')
            award_para.runs[0].font.size = Pt(11)
            award_para.paragraph_format.left_indent = Inches(0.25)

    # Languages
    if resume_model.languages:
        add_section_heading('Languages')
        lang_para = doc.add_paragraph(', '.join(resume_model.languages))
        lang_para.runs[0].font.size = Pt(11)

    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    return docx_bytes.getvalue()


def save_resume_files(
    resume_model: ResumeModel,
    output_dir: Path,
    base_filename: str = "resume"
) -> dict:
    """
    Save resume in multiple formats to specified directory.

    Args:
        resume_model: The resume data to save
        output_dir: Directory to save files to
        base_filename: Base name for the files (without extension)

    Returns:
        Dictionary with paths to saved files:
        {
            'pdf': Path or None,
            'docx': Path,
            'html': Path,
            'markdown': Path
        }
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    results = {
        'pdf': None,
        'docx': None,
        'html': None,
        'markdown': None
    }

    # Save DOCX
    docx_path = output_dir / f"{base_filename}.docx"
    docx_bytes = generate_docx(resume_model)
    docx_path.write_bytes(docx_bytes)
    results['docx'] = docx_path

    # Save HTML
    html_path = output_dir / f"{base_filename}.html"
    html_content = generate_html(resume_model)
    html_path.write_text(html_content, encoding='utf-8')
    results['html'] = html_path

    # Save Markdown
    markdown_path = output_dir / f"{base_filename}.md"
    markdown_content = resume_model.to_markdown()
    markdown_path.write_text(markdown_content, encoding='utf-8')
    results['markdown'] = markdown_path

    # Try to save PDF (optional, requires weasyprint)
    try:
        pdf_bytes = generate_pdf(resume_model)
        pdf_path = output_dir / f"{base_filename}.pdf"
        pdf_path.write_bytes(pdf_bytes)
        results['pdf'] = pdf_path
    except ImportError:
        # PDF generation not available, skip
        pass

    return results
