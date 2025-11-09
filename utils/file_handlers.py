"""File handling utilities for various resume formats."""

import json
from pathlib import Path
from typing import Tuple, Optional
import io

try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

from config.settings import (
    SUPPORTED_RESUME_FORMATS,
    MAX_FILE_SIZE_BYTES
)


def detect_file_type(file) -> str:
    """
    Detect the file type from uploaded file.

    Args:
        file: Streamlit uploaded file object

    Returns:
        File extension (e.g., 'pdf', 'docx', 'txt')
    """
    if hasattr(file, 'name'):
        return Path(file.name).suffix.lower().replace('.', '')
    return 'unknown'


def validate_file_size(file) -> Tuple[bool, str]:
    """
    Validate that file size is within acceptable limits.

    Args:
        file: Streamlit uploaded file object

    Returns:
        Tuple of (is_valid, error_message)
    """
    try:
        file.seek(0, 2)  # Seek to end
        size = file.tell()
        file.seek(0)  # Reset to beginning

        if size > MAX_FILE_SIZE_BYTES:
            size_mb = size / (1024 * 1024)
            max_mb = MAX_FILE_SIZE_BYTES / (1024 * 1024)
            return False, f"File size ({size_mb:.2f} MB) exceeds maximum allowed size ({max_mb} MB)"

        return True, ""
    except Exception as e:
        return False, f"Error checking file size: {str(e)}"


def extract_text_from_pdf(file) -> Tuple[bool, str, str]:
    """
    Extract text from PDF file using pdfplumber (primary) and PyPDF2 (fallback).

    Args:
        file: Streamlit uploaded file object

    Returns:
        Tuple of (success, text_content, error_message)
    """
    try:
        # Try pdfplumber first (better formatting preservation)
        if pdfplumber:
            try:
                file.seek(0)
                pdf = pdfplumber.open(io.BytesIO(file.read()))
                text_parts = []

                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

                pdf.close()

                if text_parts:
                    return True, '\n\n'.join(text_parts), ""
            except Exception as e:
                # Fall through to PyPDF2
                pass

        # Fallback to PyPDF2
        if PyPDF2:
            try:
                file.seek(0)
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
                text_parts = []

                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

                if text_parts:
                    return True, '\n\n'.join(text_parts), ""
                else:
                    return False, "", "Could not extract text from PDF. The file may be image-based or corrupted."
            except Exception as e:
                return False, "", f"Error reading PDF: {str(e)}"

        return False, "", "PDF processing libraries not available"

    except Exception as e:
        return False, "", f"Unexpected error processing PDF: {str(e)}"


def extract_text_from_docx(file) -> Tuple[bool, str, str]:
    """
    Extract text from DOCX file.

    Args:
        file: Streamlit uploaded file object

    Returns:
        Tuple of (success, text_content, error_message)
    """
    try:
        if not Document:
            return False, "", "python-docx library not available"

        file.seek(0)
        doc = Document(io.BytesIO(file.read()))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))

        if text_parts:
            return True, '\n'.join(text_parts), ""
        else:
            return False, "", "No text content found in DOCX file"

    except Exception as e:
        return False, "", f"Error reading DOCX file: {str(e)}"


def extract_text_from_txt(file) -> Tuple[bool, str, str]:
    """
    Extract text from plain text file.

    Args:
        file: Streamlit uploaded file object

    Returns:
        Tuple of (success, text_content, error_message)
    """
    try:
        file.seek(0)
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                text = file.read().decode(encoding)
                if text.strip():
                    return True, text, ""
            except UnicodeDecodeError:
                file.seek(0)
                continue

        return False, "", "Could not decode text file with common encodings"

    except Exception as e:
        return False, "", f"Error reading text file: {str(e)}"


def extract_text_from_md(file) -> Tuple[bool, str, str]:
    """
    Extract text from Markdown file.

    Args:
        file: Streamlit uploaded file object

    Returns:
        Tuple of (success, text_content, error_message)
    """
    # Markdown files are essentially text files
    return extract_text_from_txt(file)


def extract_text_from_json(file) -> Tuple[bool, str, str]:
    """
    Extract text from JSON file.
    Supports various JSON resume formats (JSON Resume, LinkedIn export, etc.)

    Args:
        file: Streamlit uploaded file object

    Returns:
        Tuple of (success, text_content, error_message)
    """
    try:
        file.seek(0)
        data = json.loads(file.read().decode('utf-8'))

        # Convert JSON to readable text format
        text_parts = []

        def extract_from_dict(d, prefix=""):
            """Recursively extract text from nested dictionary."""
            if isinstance(d, dict):
                for key, value in d.items():
                    if isinstance(value, (dict, list)):
                        extract_from_dict(value, f"{prefix}{key}: ")
                    elif value and str(value).strip():
                        text_parts.append(f"{prefix}{key}: {value}")
            elif isinstance(d, list):
                for item in d:
                    extract_from_dict(item, prefix)

        # Special handling for common JSON resume formats
        if 'basics' in data:  # JSON Resume format
            text_parts.append("=== BASIC INFORMATION ===")
            if 'name' in data['basics']:
                text_parts.append(f"Name: {data['basics']['name']}")
            if 'email' in data['basics']:
                text_parts.append(f"Email: {data['basics']['email']}")
            if 'phone' in data['basics']:
                text_parts.append(f"Phone: {data['basics']['phone']}")
            if 'summary' in data['basics']:
                text_parts.append(f"Summary: {data['basics']['summary']}")

            if 'work' in data:
                text_parts.append("\n=== WORK EXPERIENCE ===")
                for work in data['work']:
                    text_parts.append(f"\n{work.get('position', 'Position')} at {work.get('company', 'Company')}")
                    if 'summary' in work:
                        text_parts.append(work['summary'])

            if 'education' in data:
                text_parts.append("\n=== EDUCATION ===")
                for edu in data['education']:
                    text_parts.append(f"\n{edu.get('studyType', '')} in {edu.get('area', '')} from {edu.get('institution', '')}")

            if 'skills' in data:
                text_parts.append("\n=== SKILLS ===")
                for skill in data['skills']:
                    text_parts.append(f"{skill.get('name', '')}: {', '.join(skill.get('keywords', []))}")
        else:
            # Generic JSON extraction
            extract_from_dict(data)

        if text_parts:
            return True, '\n'.join(text_parts), ""
        else:
            return False, "", "No text content found in JSON file"

    except json.JSONDecodeError as e:
        return False, "", f"Invalid JSON format: {str(e)}"
    except Exception as e:
        return False, "", f"Error reading JSON file: {str(e)}"


def extract_text_from_file(file) -> Tuple[bool, str, str]:
    """
    Main function to extract text from uploaded file based on file type.

    Args:
        file: Streamlit uploaded file object

    Returns:
        Tuple of (success, text_content, error_message)
    """
    # Validate file size first
    size_valid, size_error = validate_file_size(file)
    if not size_valid:
        return False, "", size_error

    # Detect file type
    file_type = detect_file_type(file)

    # Route to appropriate extractor
    extractors = {
        'pdf': extract_text_from_pdf,
        'docx': extract_text_from_docx,
        'txt': extract_text_from_txt,
        'md': extract_text_from_md,
        'json': extract_text_from_json
    }

    extractor = extractors.get(file_type)
    if not extractor:
        supported = ', '.join(SUPPORTED_RESUME_FORMATS)
        return False, "", f"Unsupported file format: .{file_type}. Supported formats: {supported}"

    return extractor(file)


def save_file_to_folder(content: str, filename: str, folder: str) -> Tuple[bool, str]:
    """
    Save content to a file in the specified folder.

    Args:
        content: Text content to save
        filename: Name of the file
        folder: Destination folder path

    Returns:
        Tuple of (success, file_path or error_message)
    """
    try:
        folder_path = Path(folder)
        folder_path.mkdir(parents=True, exist_ok=True)

        file_path = folder_path / filename
        file_path.write_text(content, encoding='utf-8')

        return True, str(file_path)
    except Exception as e:
        return False, f"Error saving file: {str(e)}"
